"""
doc_parser.py - 统一文档解析器
支持 PDF / Word / Excel 三种格式，自动提取文本、内嵌图片和表格。
内嵌图片将保存到本地 backend/data/knowledge/images/ 目录。
"""
import os
import re
import hashlib
import uuid
from dataclasses import dataclass, field
from typing import List, Optional
from pathlib import Path


@dataclass
class ExtractedImage:
    """从文档中提取的图片"""
    image_id: str           # 唯一标识
    page_num: int           # 所在页码
    local_path: str         # 本地保存路径
    description: str = ""   # Gemini Vision 描述（后续填充）
    section_heading: str = "" # [NEW] 所属章节标题，用于分块定位


@dataclass
class Section:
    """文档段落"""
    heading: str            # 标题（如果有）
    content: str            # 段落内容
    page_num: int           # 所在页码
    level: int = 0          # 标题级别 (0=正文, 1=H1, 2=H2, ...)


@dataclass
class ParsedDocument:
    """解析后的文档统一格式"""
    filename: str
    file_type: str          # pdf / docx / xlsx
    full_text: str = ""     # 完整纯文本
    sections: List[Section] = field(default_factory=list)
    images: List[ExtractedImage] = field(default_factory=list)
    tables: List[str] = field(default_factory=list)
    total_pages: int = 0
    metadata: dict = field(default_factory=dict)  # 文件级元数据（source、sheets 等）
    summary: str = ""       # [NEW] LLM 生成的文档摘要
    tags: List[str] = field(default_factory=list) # [NEW] LLM 提取的语义标签

    async def enrich_with_llm(self):
        """
        [ENHANCED] 调用 LLM 提取文档摘要、标签，并智能分流处理文档内图片 (OCR/Vision)
        """
        from app.core import llm_config
        from openai import AsyncOpenAI
        from app.rag.ingestion.vision_describer import describe_image_with_vision
        from app.tools.ocr_engine import get_ocr_processor
        import asyncio
        import json

        print(f"[DocParser] 🚀 enrich_with_llm() 开始执行...")
        print(f"[DocParser] 📝 正在获取 LLM 配置...")
        base_url, api_key = await llm_config.get_llm_config()
        print(f"[DocParser] ✅ 获取配置成功: base_url={base_url is not None}, api_key={api_key is not None}")
        if not api_key:
            print("[DocParser] ⚠️ 没有 API Key，跳过增强")
            return
            
        client = AsyncOpenAI(api_key=api_key, base_url=base_url)
        
        # 1. 文本语义增强 (摘要与标签)
        sample_text = self.full_text[:4000]
        prompt = f"""
        你是一个精通文档管理与检索优化的专家。请分析以下文档内容，并提供：
        1. 摘要：用不超过 150 字概括文档的核心信息，用于辅助检索时的意图理解。
        2. 标签：根据文档内容实事求是地提取 5-8 个关键词标签。
        文档内容片段：
        {sample_text}
        请以 JSON 格式返回：{{"summary": "...", "tags": ["tag1", ...]}}
        """
        
        try:
            print(f"[DocParser] 🧠 正在为文档 '{self.filename}' 生成语义增强信息...")
            text_task = client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],
                response_format={ "type": "json_object" }
            )
            
            # 2. 图片智能处理 (并发分流)
            image_tasks = []
            ocr_processor = get_ocr_processor()

            async def process_single_image(img: ExtractedImage):
                try:
                    # 1. 准备图片字节流用于 Vision
                    with open(img.local_path, "rb") as f:
                        img_bytes = f.read()
                    
                    # 2. 异步并行执行：Vision 描述 + 本地 OCR 识别
                    # [Hybrid] 综合 Vision 的“结构感”与 PaddleOCR 的“字准度”
                    vision_task = describe_image_with_vision(image_buffer=img_bytes, title=f"所处文档: {self.filename}", mode="describe")
                    
                    loop = asyncio.get_event_loop()
                    ocr_task = loop.run_in_executor(None, ocr_processor.ocr_image_from_file, img.local_path)
                    
                    vision_res, ocr_text = await asyncio.gather(vision_task, ocr_task)
                    
                    # 3. 结果深度融合
                    is_table = vision_res.get("is_table", False)
                    vision_desc = vision_res.get("description", "").strip()
                    
                    if is_table:
                        print(f"  └── 🔍 识别到表格/数据图: {img.image_id}, 正在全量融合 Vision 结构化 Markdown 与 OCR 文本...")
                    
                    # 组装最终描述：Vision 描述在前（供阅读），OCR 原文在后（供检索）
                    merged_content = vision_desc
                    if ocr_text and ocr_text.strip():
                        # 为避免内容重复导致的干扰，仅在 OCR 有实质内容时添加标签
                        merged_content += f"\n\n---\n[OCR 无损原文 (用于精确检索)]\n{ocr_text.strip()}"
                    
                    img.description = merged_content
                    
                except Exception as e:
                    print(f"  └── ❌ 图片处理失败 ({img.image_id}): {e}")

            if self.images:
                print(f"[DocParser] 📸 正在处理 {len(self.images)} 张提取到的图片...")
                image_tasks = [process_single_image(img) for img in self.images[:10]] # 限制前 10 张，避免超长等待

            # 执行所有并发任务
            results = await asyncio.gather(text_task, *image_tasks)
            
            #  text_task 已经完成，通过 result() 获取第一项，避免 RuntimeError
            text_resp = results[0]
            result = json.loads(text_resp.choices[0].message.content)
            self.summary = result.get("summary", "")
            self.tags = result.get("tags", [])
            print(f"[DocParser] ✅ 全流程增强完成: {len(self.tags)} 标签, {len(self.images)} 图片已处理")
        except Exception as e:
            print(f"[DocParser] ⚠️ 增强流程异常: {e}")

    def to_llama_document(self) -> "Document":
        """
        转换为 LlamaIndex 的 Document 对象
        """
        from llama_index.core import Document
        
        # 构建基础元数据
        metadata = {
            "filename": self.filename,
            "file_type": self.file_type,
            "total_pages": self.total_pages,
            "doc_summary": self.summary,
            "doc_tags": ",".join(self.tags),
            **self.metadata
        }
        
        return Document(
            text=self.full_text,
            metadata=metadata,
            excluded_llm_metadata_keys=["file_type", "total_pages"],
            excluded_embed_metadata_keys=["file_type", "total_pages", "doc_summary", "doc_tags"]
        )


# ─── 图片保存目录 ──────────────────────────────────────────────────────────────

def _get_images_dir() -> str:
    """获取图片存储目录"""
    current_dir = os.path.dirname(os.path.abspath(__file__))
    images_dir = os.path.abspath(os.path.join(current_dir, "..", "..", "..", "data", "knowledge", "images"))
    os.makedirs(images_dir, exist_ok=True)
    return images_dir


def _rows_to_markdown_table(rows: List[List[str]]) -> str:
    """将二维列表转换为标准 Markdown 表格"""
    if not rows:
        return ""
    
    # 过滤掉全空行
    rows = [[str(c).replace('\n', ' ').strip() for c in r] for r in rows if any(str(c).strip() for c in r)]
    if not rows:
        return ""

    # 计算列数
    num_cols = max(len(r) for r in rows)
    
    md_rows = []
    # 表头处理
    header = rows[0]
    # 填充对齐
    if len(header) < num_cols:
        header.extend([""] * (num_cols - len(header)))
    md_rows.append("| " + " | ".join(header) + " |")
    
    # 分隔符
    md_rows.append("| " + " | ".join(["---"] * num_cols) + " |")
    
    # 数据行
    for row in rows[1:]:
        if len(row) < num_cols:
            row.extend([""] * (num_cols - len(row)))
        md_rows.append("| " + " | ".join(row) + " |")
        
    return "\n".join(md_rows)


# ─── PDF 解析 ──────────────────────────────────────────────────────────────────

def parse_pdf(file_path: str) -> ParsedDocument:
    """解析 PDF 文件：提取文本 + 内嵌图片"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")

    doc = fitz.open(file_path)
    filename = os.path.basename(file_path)
    images_dir = _get_images_dir()

    sections = []
    images = []
    full_text_parts = []

    for page_num in range(len(doc)):
        page = doc[page_num]

        # 提取文本
        text = page.get_text("text").strip()
        if text:
            full_text_parts.append(text)
            # 尝试识别标题（大字号或粗体的首行）
            blocks = page.get_text("blocks")
            current_heading = ""
            current_content = []

            for block in blocks:
                if len(block) >= 5:
                    block_text = block[4].strip() if isinstance(block[4], str) else ""
                    if block_text:
                            # 增强型标题识别：不仅看位置，还看样式
                            font_size = block[0] if isinstance(block[0], (int, float)) else 0
                            # 如果字号明显大于普通段落 (通常 > 12) 或有粗体标记，认定为标题
                            is_strong_heading = font_size > 14 or (len(block_text) < 40 and block[1] < page.rect.height * 0.2)
                            
                            if is_strong_heading:
                                if current_content:
                                    sections.append(Section(
                                        heading=current_heading,
                                        content="\n".join(current_content),
                                        page_num=page_num + 1,
                                        level=1 if current_heading else 0
                                    ))
                                    current_content = []
                                current_heading = block_text
                            else:
                                current_content.append(block_text)

            # 收尾
            if current_content:
                sections.append(Section(
                    heading=current_heading,
                    content="\n".join(current_content),
                    page_num=page_num + 1,
                    level=1 if current_heading else 0
                ))

        # 提取图片
        try:
            image_list = page.get_images(full=True)
            for img_idx, img_info in enumerate(image_list):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if base_image and base_image.get("image"):
                        img_ext = base_image.get("ext", "png")
                        # 改进：使用文件名哈希而非全路径，避免临时路径导致的失败
                        file_ref = filename or "unknown_doc"
                        img_id = f"{hashlib.md5(file_ref.encode()).hexdigest()[:8]}_p{page_num+1}_i{img_idx}"
                        img_path = os.path.join(images_dir, f"{img_id}.{img_ext}")

                        if not os.path.exists(img_path):
                            with open(img_path, "wb") as f:
                                f.write(base_image["image"])

                        images.append(ExtractedImage(
                            image_id=img_id,
                            page_num=page_num + 1,
                            local_path=img_path,
                            section_heading=current_heading # [PDF 增强] 关联本页识别到的最后一个标题
                        ))
                except Exception as e:
                    print(f"[DocParser] 提取单张图片失败 (page {page_num+1}, xref {xref}): {e}")
        except Exception as e:
            print(f"[DocParser] 获取页面图片列表失败 (page {page_num+1}): {e}")

    num_pages = len(doc)
    doc.close()

    # [增强] 解析后聚合：将过短且属于同一标题的 Section 进行合并，减少碎片
    merged_sections = []
    if sections:
        curr = sections[0]
        for next_sec in sections[1:]:
            # 如果下一段没有新标题，或者内容极短（<100字），尝试合并到上一段
            if not next_sec.heading or (len(next_sec.content) < 150 and next_sec.heading == curr.heading):
                curr.content += "\n" + next_sec.content
            else:
                merged_sections.append(curr)
                curr = next_sec
        merged_sections.append(curr)

    return ParsedDocument(
        filename=filename,
        file_type="pdf",
        full_text="\n\n".join(full_text_parts),
        sections=merged_sections,
        images=images,
        total_pages=num_pages,
        metadata={"source": filename}
    )


# ─── Word 解析 ─────────────────────────────────────────────────────────────────

def iter_block_items(parent, doc_obj=None):
    """
    深度递归遍历 Word 文档中的所有内容块（段落和表格），保持物理顺序。
    支持从 Document, Table, Row, Cell 全平面展开。
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table, _Row, _Cell
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.document import Document

    if doc_obj is None and isinstance(parent, Document):
        doc_obj = parent

    # 获取底层 xml 元素
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, Table):
        for row in parent.rows:
            yield from iter_block_items(row, doc_obj)
        return
    elif isinstance(parent, _Row):
        for cell in parent.cells:
            yield from iter_block_items(cell, doc_obj)
        return
    else:
        return

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent if not isinstance(parent, Document) else doc_obj)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent if not isinstance(parent, Document) else doc_obj)


def parse_docx(file_path: str) -> ParsedDocument:
    """解析 Word 文件：提取段落 + 行内图片 + 表格"""
    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
    except ImportError:
        raise ImportError("请安装 python-docx: pip install python-docx")

    from docx.text.paragraph import Paragraph
    from docx.table import Table

    doc = Document(file_path)
    filename = os.path.basename(file_path)
    images_dir = _get_images_dir()

    sections = []
    images = []
    tables = []
    full_text_parts = []
    current_heading = ""
    current_content = []

    # [NEW] 核心命名空间字典，用于 XPath 查找内嵌资源
    DOCX_NSMAP = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'v': 'urn:schemas-microsoft-com:vml'
    }

    def _extract_images_from_element(element, current_heading, current_page):
        """内部辅助函数：从任意 XML 元素中提取图片"""
        from lxml import etree
        # [FIX] python-docx 的 BaseOxmlElement 不直接支持 namespaces 参数，改用 lxml 原生方法
        xpath_str = './/a:blip/@r:embed | .//v:imagedata/@r:id | .//pic:blipFill/a:blip/@r:embed'
        try:
            rIds = element.xpath(xpath_str, namespaces=DOCX_NSMAP)
        except TypeError:
            # 兜底：如果直接调用失败，尝试使用 lxml.etree 包装
            try:
                finder = etree.XPath(xpath_str, namespaces=DOCX_NSMAP)
                rIds = finder(element)
            except Exception as e:
                print(f"[DocParser] XPath 搜索严重失败: {e}")
                rIds = []

        for rId in rIds:
            try:
                rel = doc.part.rels[rId]
                if "image" in rel.reltype:
                    img_data = rel.target_part.blob
                    img_ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "png"
                    img_id = f"{hashlib.md5(file_path.encode()).hexdigest()[:8]}_img_{uuid.uuid4().hex[:6]}"
                    img_path = os.path.join(images_dir, f"{img_id}.{img_ext}")
                    
                    with open(img_path, "wb") as f:
                        f.write(img_data)
                    
                    images.append(ExtractedImage(
                        image_id=img_id,
                        page_num=current_page,
                        local_path=img_path,
                        section_heading=current_heading
                    ))
                    print(f"[DocParser] 📸 提取到图片: {img_id}, 关联标题: {current_heading[:20]}")
            except Exception as e:
                pass # 忽略单个提取失败

    # 标题级别映射
    heading_priority = {
        'Title': 1, 'Subtitle': 2, 'Heading 1': 1, 'Heading 2': 2, 'Heading 3': 3,
        'Heading 4': 4, 'Heading 5': 5, 'Heading 6': 6
    }

    # 使用块迭代器，确保表格不丢失且顺序正确
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            
            # --- 使用辅助函数处理行内图片 ---
            _extract_images_from_element(block._p, current_heading, 1)

            if not text:
                continue

            full_text_parts.append(text)
            style_name = block.style.name if block.style else "Normal"
            
            # 正格正则匹配：支持 Part/一、/第一章等常见结构
            is_heading_style = style_name in heading_priority or \
                        ("Heading" in style_name and len(text) < 100) or \
                        (style_name == "Title")
            
            is_heading_pattern = bool(re.search(r'(Part\s*\d+|第[一二三四五六七八九十]+\s*[章节]|第\s*\d+\s*[章节]|[一二三四五六七八九十]、)', text, re.I)) and len(text) < 100
            
            is_heading = is_heading_style or is_heading_pattern

            if is_heading:
                print(f"[DocParser] 📍 识别到标题: '{text[:30]}...'")
                if current_content:
                    sections.append(Section(
                        heading=current_heading,
                        content="\n".join(current_content),
                        page_num=1,
                        level=1 if current_heading else 0
                    ))
                    current_content = []
                current_heading = text
            else:
                current_content.append(text)
        
        elif isinstance(block, Table):
            # 处理表格：在此改为直接展开提取表格内的文本以防丢失标题
            # 如果表格非常大，我们通过递归其单元格来处理
            for row in block.rows:
                row_cells_content = []
                for cell in row.cells:
                    # 获取单元格内的所有文本块
                    for cell_block in iter_block_items(cell, doc):
                        if isinstance(cell_block, Paragraph):
                            cell_text = cell_block.text.strip()
                            if cell_text:
                                # [关键] 检查单元格内是否包含新章节标题
                                if bool(re.search(r'(Part\s*\d+|第[一二三四五六七八九十]+\s*[章节]|第\s*\d+\s*[章节]|[一二三四五六七八九十]、)', cell_text, re.I)) and len(cell_text) < 100:
                                    print(f"[DocParser] 🚨 在表格单元格中发现标题: '{cell_text[:30]}...'")
                                    if current_content:
                                        sections.append(Section(heading=current_heading, content="\n".join(current_content), page_num=1))
                                        current_content = []
                                    current_heading = cell_text
                                else:
                                    current_content.append(cell_text)
                                    full_text_parts.append(cell_text)
                            
                            # [FIX] 表格单元格内图片提取
                            _extract_images_from_element(cell_block._p, current_heading, 1)
                    # 记录单元格内容用于表格 MD (可选，先保文字)
                    row_cells_content.append(cell.text.strip())
                # 即使不生成 MD，也要记录表格内容
                # tables.append(" | ".join(row_cells_content))

    # 收尾
    if current_content:
        sections.append(Section(
            heading=current_heading,
            content="\n".join(current_content),
            page_num=1,
            level=1 if current_heading else 0
        ))

    # ── [终极兜底] 强制 100% 文本留存 ──
    # 计算期望的总原文（包含所有段落和表格文字）
    def get_all_raw_text(doc_obj):
        texts = []
        # 使用原生 paragraphs
        for p in doc_obj.paragraphs:
            if p.text.strip(): texts.append(p.text.strip())
        # 使用原生 tables (递归)
        def _get_table_text(table_obj):
            t_texts = []
            for row in table_obj.rows:
                for cell in row.cells:
                    if cell.tables:
                        for sub_t in cell.tables:
                            t_texts.extend(_get_table_text(sub_t))
                    t_texts.append(cell.text.strip())
            return t_texts
        for t in doc_obj.tables:
            texts.extend(_get_table_text(t))
        return "\n".join(texts)

    raw_text_stream = get_all_raw_text(doc).strip()
    current_total_content = "\n".join([s.content for s in sections]).strip()
    
    # 只要差额超过 100 字符，就判定为丢失，触发强力补齐
    if len(raw_text_stream) > len(current_total_content) + 100:
        print(f"[DocParser] 🚨 严重内容丢失警报！原文 {len(raw_text_stream)} 字，提取 {len(current_total_content)} 字。正在强行补齐...")
        # 补丁方案：在末尾追加所有未在 sections 中出现的长片段，或者直接全量重排
        # 为保险起见，这里直接使用全量流式重排，因为用户要求“一个字不能少”
        sections = []
        current_content = []
        all_lines = raw_text_stream.split("\n")
        for line in all_lines:
            if not line.strip(): continue
            current_content.append(line)
            if sum(len(c) for c in current_content) > 1000:
                sections.append(Section(heading="[全量留存分块]", content="\n".join(current_content), page_num=1))
                current_content = []
        if current_content:
            sections.append(Section(heading="[全量留存分块]", content="\n".join(current_content), page_num=1))
    
    # 提取完成后不再从 rels 全量扫一遍，因为 Paragraph 已经按顺序处理了 (除非有浮动图片)
    # 作为一个兜底，扫描一下不在已追踪列表里的图片是有必要的，但为了保持顺序，Paragraph 优先。
    
    return ParsedDocument(
        filename=filename,
        file_type="docx",
        full_text="\n\n".join(full_text_parts),
        sections=sections,
        images=images,
        tables=tables,
        total_pages=1,
        metadata={"source": file_path}
    )


# ─── Excel 解析 ────────────────────────────────────────────────────────────────

def parse_xlsx(file_path: str) -> ParsedDocument:
    """解析 Excel 文件：每个 Sheet 转化为结构化文本"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("请安装 openpyxl: pip install openpyxl")

    wb = load_workbook(file_path, data_only=True)
    filename = os.path.basename(file_path)

    sections = []
    tables = []
    full_text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            rows.append(cells)

        if rows:
            table_md = _rows_to_markdown_table(rows)
            if table_md:
                table_text = f"\n[Sheet: {sheet_name}]\n{table_md}\n"
                tables.append(table_text)
                full_text_parts.append(table_text)

                sections.append(Section(
                    heading=f"Sheet: {sheet_name}",
                    content=table_text,
                    page_num=wb.sheetnames.index(sheet_name) + 1,
                    level=1
                ))

    wb.close()

    return ParsedDocument(
        filename=filename,
        file_type="xlsx",
        full_text="\n\n".join(full_text_parts),
        sections=sections,
        tables=tables,
        total_pages=len(wb.sheetnames),
        metadata={"source": file_path, "sheets": wb.sheetnames}
    )


# ─── 统一入口 ──────────────────────────────────────────────────────────────────

SUPPORTED_TYPES = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".xlsx": parse_xlsx,
    ".doc": parse_docx,   # 兼容：尝试用 python-docx 打开
    ".xls": parse_xlsx,   # 兼容：尝试用 openpyxl 打开
}


def parse_document(file_path: str) -> ParsedDocument:
    """
    统一文档解析入口
    根据文件扩展名自动选择解析器

    Args:
        file_path: 文件绝对路径

    Returns:
        ParsedDocument 对象

    Raises:
        ValueError: 不支持的文件格式
    """
    ext = os.path.splitext(file_path)[1].lower()
    parser = SUPPORTED_TYPES.get(ext)

    if parser is None:
        supported = ", ".join(SUPPORTED_TYPES.keys())
        raise ValueError(f"不支持的文件格式: {ext}。支持: {supported}")

    print(f"[DocParser] 解析文件: {os.path.basename(file_path)} (类型: {ext})")
    result = parser(file_path)
    print(f"[DocParser] ✅ 解析完成: {len(result.sections)} 段落, {len(result.images)} 图片, {len(result.tables)} 表格")
    return result
